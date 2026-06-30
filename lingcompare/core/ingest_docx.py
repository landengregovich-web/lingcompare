"""Parse Word documents (.docx) into Corpus objects.

Supports two layouts:
  1. Table layout (preferred): each row in the first table is one entry.
     Column detection follows the same alias logic as ingest_xlsx.
  2. Paragraph layout (fallback): lines matching "gloss TAB ipa" or
     "gloss COMMA ipa" are extracted when no table is present.

Multiple tables in the document are all processed and combined.
"""

from __future__ import annotations

import docx

from .ingest_wordlist import (
    ValidationError,
    _validate_and_tokenize_ipa,
    _parse_ipa_tokens,
    _parse_breakdown,
    _looks_like_header,
)
from .ingest_xlsx import _find_columns, _GLOSS_ALIASES, _IPA_ALIASES
from .schema import Corpus, Word


def _rows_from_table(table) -> list[list[str]]:
    """Extract text rows from a python-docx Table object."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        # Merged cells can appear as duplicates — deduplicate consecutive identical cells
        deduped: list[str] = []
        for c in cells:
            if not deduped or c != deduped[-1]:
                deduped.append(c)
        rows.append(deduped)
    return rows


def _rows_from_paragraphs(doc) -> list[list[str]]:
    """Fallback: split paragraphs on tab or comma to extract (gloss, IPA) pairs."""
    rows = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if "\t" in text:
            parts = [p.strip() for p in text.split("\t", maxsplit=2)]
        elif "," in text:
            parts = [p.strip() for p in text.split(",", maxsplit=2)]
        else:
            continue
        if len(parts) >= 2:
            rows.append(parts)
    return rows


def _parse_rows(
    rows: list[list[str]],
    language_name: str,
    row_offset: int = 0,
) -> tuple[list[Word], list[ValidationError]]:
    """Convert a list of cell-rows into Word objects."""
    errors: list[ValidationError] = []
    words: list[Word] = []

    if not rows:
        return words, errors

    # Header detection
    gloss_col = ipa_col = bd_col = None
    data_start = 0

    if len(rows) >= 2:
        first = rows[0]
        second_cell = first[1] if len(first) > 1 else ""
        is_header = _looks_like_header(second_cell) or (
            first[0].strip().lower() in _GLOSS_ALIASES or
            second_cell.strip().lower() in _IPA_ALIASES
        )
        if is_header:
            gloss_col, ipa_col, bd_col = _find_columns(first)
            data_start = 1

    if gloss_col is None:
        gloss_col = 0
    if ipa_col is None:
        ipa_col = 1

    for i, row in enumerate(rows[data_start:], start=data_start + 1 + row_offset):
        if not any(row):
            continue
        if len(row) <= max(gloss_col, ipa_col):
            errors.append(ValidationError(
                row=i, column="IPA_form",
                message=f"Row has only {len(row)} column(s); need at least {max(gloss_col, ipa_col) + 1}.",
            ))
            continue

        gloss = row[gloss_col].strip()
        ipa_raw = row[ipa_col].strip()
        breakdown_raw = row[bd_col].strip() if bd_col is not None and bd_col < len(row) else ""

        if not gloss:
            errors.append(ValidationError(row=i, column="gloss", message="Gloss is empty."))
            continue
        if not ipa_raw:
            errors.append(ValidationError(row=i, column="IPA_form", message="IPA form is empty."))
            continue

        valid_tokens, tok_errors = _validate_and_tokenize_ipa(ipa_raw, i, "IPA_form")
        errors.extend(tok_errors)
        if tok_errors:
            continue

        if not valid_tokens:
            errors.append(ValidationError(
                row=i, column="IPA_form",
                message=f"No recognised IPA segments in {ipa_raw!r}.",
            ))
            continue

        segments, seg_errors = _parse_ipa_tokens(valid_tokens, i, "IPA_form")
        errors.extend(seg_errors)
        if seg_errors:
            continue

        morphemes = []
        if breakdown_raw:
            morphemes, bd_errors = _parse_breakdown(ipa_raw, breakdown_raw, i)
            errors.extend(bd_errors)
            if bd_errors:
                morphemes = []

        words.append(Word(gloss=gloss, segments=segments, morphemes=morphemes))

    return words, errors


def parse(
    file_bytes: bytes,
    language_name: str = "Unknown",
) -> tuple[Corpus, list[ValidationError]]:
    """Parse a Word document (.docx) from raw bytes.

    Args:
        file_bytes: Raw .docx file bytes.
        language_name: Assigned to the resulting Corpus.

    Returns:
        (corpus, errors) — errors is empty on full success.
    """
    import io as _io
    errors: list[ValidationError] = []
    words: list[Word] = []

    try:
        doc = docx.Document(_io.BytesIO(file_bytes))
    except Exception as exc:
        errors.append(ValidationError(
            row=0, column="file",
            message=f"Could not open Word document: {exc}",
        ))
        return Corpus(language_name=language_name), errors

    if doc.tables:
        row_offset = 0
        for table in doc.tables:
            rows = _rows_from_table(table)
            w, e = _parse_rows(rows, language_name, row_offset=row_offset)
            words.extend(w)
            errors.extend(e)
            row_offset += len(rows)
    else:
        rows = _rows_from_paragraphs(doc)
        words, errors = _parse_rows(rows, language_name)

    return Corpus(language_name=language_name, words=words), errors
