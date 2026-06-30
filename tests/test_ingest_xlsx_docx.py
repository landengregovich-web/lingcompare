"""Tests for ingest_xlsx.py and ingest_docx.py."""

import io
import pytest
import openpyxl
import docx

from lingcompare.core.ingest_xlsx import parse as parse_xlsx
from lingcompare.core.ingest_docx import parse as parse_docx


# ---------------------------------------------------------------------------
# Helpers: create in-memory Excel / Word files
# ---------------------------------------------------------------------------

def make_xlsx(rows: list[list], sheet_name: str = "Sheet1") -> bytes:
    """Create an in-memory .xlsx with one sheet containing `rows`."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = sheet_name
    for row in rows:
        ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def make_docx_table(rows: list[list[str]]) -> bytes:
    """Create an in-memory .docx with one table containing `rows`."""
    doc = docx.Document()
    if not rows:
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            table.rows[ri].cells[ci].text = str(cell_text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_docx_paragraphs(lines: list[str]) -> bytes:
    """Create a .docx with one paragraph per line (no table)."""
    doc = docx.Document()
    for line in lines:
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Excel tests
# ---------------------------------------------------------------------------

class TestParseXlsx:

    def test_basic_two_column_no_header(self):
        data = make_xlsx([["water", "aɣwa"], ["fire", "fweɣo"]])
        corpus, errors = parse_xlsx(data)
        assert errors == []
        assert len(corpus.words) == 2
        assert corpus.words[0].gloss == "water"
        assert corpus.words[1].gloss == "fire"

    def test_with_named_header(self):
        data = make_xlsx([["gloss", "IPA"], ["water", "aɣwa"], ["fire", "fweɣo"]])
        corpus, errors = parse_xlsx(data)
        assert errors == []
        assert len(corpus.words) == 2

    def test_header_alias_word_phonetic(self):
        data = make_xlsx([["word", "phonetic"], ["water", "aɣwa"]])
        corpus, errors = parse_xlsx(data)
        assert errors == []
        assert len(corpus.words) == 1

    def test_segments_populated(self):
        data = make_xlsx([["water", "aba"]])
        corpus, errors = parse_xlsx(data)
        assert errors == []
        assert len(corpus.words[0].segments) == 3

    def test_language_name_set(self):
        data = make_xlsx([["water", "aba"]])
        corpus, _ = parse_xlsx(data, language_name="TestLang")
        assert corpus.language_name == "TestLang"

    def test_empty_gloss_error(self):
        data = make_xlsx([["", "aba"]])
        _, errors = parse_xlsx(data)
        assert any("Gloss" in e.message for e in errors)

    def test_empty_ipa_error(self):
        data = make_xlsx([["water", ""]])
        _, errors = parse_xlsx(data)
        assert any("IPA" in e.message for e in errors)

    def test_invalid_ipa_error(self):
        data = make_xlsx([["water", "aQa"]])
        _, errors = parse_xlsx(data)
        assert len(errors) >= 1
        assert any("Q" in e.message for e in errors)

    def test_skips_empty_rows(self):
        data = make_xlsx([["water", "aba"], ["", ""], ["fire", "fweɣo"]])
        corpus, errors = parse_xlsx(data)
        assert errors == []
        assert len(corpus.words) == 2

    def test_third_column_as_breakdown(self):
        data = make_xlsx([
            ["gloss", "IPA", "breakdown"],
            ["water-PL", "abas", "aba-s"],
        ])
        corpus, errors = parse_xlsx(data)
        assert errors == []
        word = corpus.words[0]
        assert len(word.morphemes) == 2

    def test_swadesh_fixture_roundtrip(self):
        """Convert the CSV fixture to Excel and parse it."""
        from pathlib import Path
        import csv
        rows_csv = list(csv.reader(
            (Path(__file__).parent / "fixtures" / "es_swadesh.csv").open(encoding="utf-8")
        ))
        data = make_xlsx(rows_csv)
        corpus, errors = parse_xlsx(data, language_name="Spanish")
        assert len(errors) == 0
        assert len(corpus.words) == 15

    def test_corrupt_file_returns_error(self):
        _, errors = parse_xlsx(b"not an xlsx file")
        assert len(errors) >= 1
        assert any("workbook" in e.message.lower() for e in errors)


# ---------------------------------------------------------------------------
# Word document tests
# ---------------------------------------------------------------------------

class TestParseDocx:

    def test_basic_table_no_header(self):
        data = make_docx_table([["water", "aɣwa"], ["fire", "fweɣo"]])
        corpus, errors = parse_docx(data)
        assert errors == []
        assert len(corpus.words) == 2

    def test_table_with_named_header(self):
        data = make_docx_table([["gloss", "IPA"], ["water", "aɣwa"]])
        corpus, errors = parse_docx(data)
        assert errors == []
        assert len(corpus.words) == 1
        assert corpus.words[0].gloss == "water"

    def test_segments_populated(self):
        data = make_docx_table([["water", "aba"]])
        corpus, errors = parse_docx(data)
        assert errors == []
        assert len(corpus.words[0].segments) == 3

    def test_language_name_set(self):
        data = make_docx_table([["water", "aba"]])
        corpus, _ = parse_docx(data, language_name="TestLang")
        assert corpus.language_name == "TestLang"

    def test_invalid_ipa_error(self):
        data = make_docx_table([["water", "aQa"]])
        _, errors = parse_docx(data)
        assert any("Q" in e.message for e in errors)

    def test_empty_gloss_error(self):
        data = make_docx_table([["", "aba"]])
        _, errors = parse_docx(data)
        assert any("Gloss" in e.message for e in errors)

    def test_paragraph_fallback_tab_separated(self):
        lines = ["water\taɣwa", "fire\tfweɣo"]
        data = make_docx_paragraphs(lines)
        corpus, errors = parse_docx(data)
        assert errors == []
        assert len(corpus.words) == 2
        assert corpus.words[0].gloss == "water"

    def test_paragraph_fallback_comma_separated(self):
        lines = ["water,aba", "fire,fweɣo"]
        data = make_docx_paragraphs(lines)
        corpus, errors = parse_docx(data)
        assert errors == []
        assert len(corpus.words) == 2

    def test_paragraph_skips_lines_without_delimiter(self):
        lines = ["My wordlist", "water\taba", "just text here", "fire\tfweɣo"]
        data = make_docx_paragraphs(lines)
        corpus, errors = parse_docx(data)
        assert errors == []
        assert len(corpus.words) == 2

    def test_multiple_tables_combined(self):
        """Two tables in one doc are both parsed."""
        doc = docx.Document()
        for entries in [
            [["water", "aba"], ["fire", "fweɣo"]],
            [["hand", "mano"], ["eye", "oxo"]],
        ]:
            table = doc.add_table(rows=len(entries), cols=2)
            for ri, row in enumerate(entries):
                for ci, text in enumerate(row):
                    table.rows[ri].cells[ci].text = text
        buf = io.BytesIO()
        doc.save(buf)
        corpus, errors = parse_docx(buf.getvalue())
        assert errors == []
        assert len(corpus.words) == 4

    def test_corrupt_file_returns_error(self):
        _, errors = parse_docx(b"not a docx file")
        assert len(errors) >= 1
        assert any("Word" in e.message for e in errors)

    def test_swadesh_fixture_roundtrip(self):
        """Round-trip the CSV fixture through a Word table."""
        from pathlib import Path
        import csv
        rows_csv = list(csv.reader(
            (Path(__file__).parent / "fixtures" / "es_swadesh.csv").open(encoding="utf-8")
        ))
        data = make_docx_table([[str(c) for c in row] for row in rows_csv])
        corpus, errors = parse_docx(data, language_name="Spanish")
        assert errors == []
        assert len(corpus.words) == 15
